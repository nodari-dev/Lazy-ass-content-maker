import random
import presets
import requests
from bs4 import BeautifulSoup as bs
from docx import Document
import os


websitesDict = {}
resultDict = {}

websitesList = []
citiesList = []

# categoryList = ['design', 'development', 'shopify', 'wordpress']
categoryList = ['development']

# TODO: better filtering 'Wordpress or woocommerce' ==> 'New category'
#   Find all keywords
# TODO: Fix mainContainer data under subFooter content
# TODO: According to heading add words to paragraph FRONT-END DEVELOPMENT


def getFile(lineName, procList):
    """
    This function opens txt file
    Then add every line to array we will work with
    :param lineName:
    :param procList:
    """
    try:
        openedFile = open(lineName, "r")
    except NameError:
        print('No file named ' + lineName)
    else:
        for line in openedFile:
            lineElement = line.strip()
            procList.append(lineElement)
    finally:
        openedFile.close()


def parseWebsites():
    """
    This function parse all websites
    1. We parse all HTML PAGE, all page in variable 'website'
    2. We run getCityName(website) to return name of city
    3. Then we run getCategory(website) to return category
    4. Then we get underHeader (subHeader)
    5. Then we get above-footer (subFooter)
    """
    for website in websitesList:
        print(website)
        # Parse all HTML page
        page = requests.get(website)
        websiteContent = bs(page.content, 'html.parser')
        # Write all changes to dictionary

        websitesDict[website] = {
            "City": getCityName(website),
            "Category": getCategory(website),
        }

        # We use these functions to get all content from pages
        # Sort and add them to websiteDict
        getSubHeaderFooter('subHeader', website, websiteContent)
        getMainContainer(website, websiteContent)
        getSubHeaderFooter('subFooter', website, websiteContent)


def getCityName(websiteProc):
    """
    This function get city name from link
    First we need to remove chars for presets.checks
    Then clean text and set correct capitalize

    :param websiteProc
    """
    websitesCity = str(websiteProc)

    # Remove checking elements
    for check, empty in presets.checks.items():
        if check in websitesCity:
            websitesCity = websitesCity.replace(check, empty)
            break

    # Clean text
    for key, value in presets.charToReplace.items():
        websitesCity = websitesCity.replace(key, value)

        # Capitalize first and last word
        websiteCityCapitalize = websitesCity.split()

        # Capitalize first word
        websiteCityCapitalize[0] = websiteCityCapitalize[0].capitalize()

        # Capitalize last word if length > 1
        if len(websiteCityCapitalize) > 1:
            websiteCityCapitalize[-1] = websiteCityCapitalize[-1].capitalize()
        websiteCity = ' '.join(websiteCityCapitalize)

    return websiteCity


def getCategory(websiteProc):
    """
    This function identifying category from each link

    :param websiteProc
    """
    for key, category in presets.categoryDict.items():
        if key in websiteProc:
            return category
            break


def getSubHeaderFooter(procKey, linkProc, websiteProcContent):
    """
    This function get heading and paragraph from div`s 'underheader' and 'above-footer'
    Then write changes into websiteDict

    :param procKey:
    :param linkProc:
    :param websiteProcContent
    """
    # Check which block we need to get
    if 'Header' in procKey:
        elements = websiteProcContent.find_all('div', class_='container_full underheader')
    else:
        elements = websiteProcContent.find_all('div', class_='container_full above-footer')
    try:
        elements[0]
        # Catch error if block does not exist
    except IndexError:
        pass
    else:
        if len(elements) != 0:
            for textEl in elements:
                if 'h2' in str(textEl):
                    # First we get <h2> text and if exist such element
                    heading = textEl.text
                    # 'n' variable is a count of found elements with same type
                    for n, element in enumerate(elements, start=0):
                        # We use procKey to create full correct key [subHeaderHeading]
                        websitesDict[linkProc][procKey + 'Heading' + str(n)] = heading
                        break
                else:
                    # First we get <p> text and if exist such element
                    paragraph = textEl.text
                    for n, element in enumerate(elements, start=0):
                        # We use procKey to create full correct key [subHeaderParagraph]
                        websitesDict[linkProc][procKey + 'Paragraph' + str(n)] = paragraph
                        break


def getMainContainer(linkProc, websiteProcContent):
    """
    This function gets all headings and paragraphs
    from <div class="main-container-landing">
    And write them into websitesDict

    :param linkProc:
    :param websiteProcContent:
    """
    elements = websiteProcContent.find_all('div', class_='main-container-landing')
    try:
        elements[0]
    except NameError:
        # Catch error if block does not exist
        print('No div with class "main-container-landing"')
    else:
        # 'n' variable is a count of found elements with same type
        for n, element in enumerate(elements, start=0):
            # here we find paragraph
            paragraph = element.find('p').text
            websitesDict[linkProc]['mainParagraph' + str(n)] = paragraph


def contentCreator(mainCity, mainCategory):
    """
    This function add random and specific content do resultDict
    1: we need to check City and Category of website we need to create
    2: select by City and Category <title> and <meta>
    3: add random content to resultDict

    NOTE: every function replace city and category to return final text
    :param mainCity:
    :param mainCategory:
    """
    #
    resultDict['city'] = mainCity
    resultDict['mainCategory'] = mainCategory
    resultDict['title'] = filterTitle(mainCity, mainCategory)
    resultDict['meta'] = filterMeta(mainCity, mainCategory)
    resultDict['headerHeading'] = filterHeaderHeading(mainCity, mainCategory)
    resultDict['headerParagraph'] = filterHeaderParagraph(mainCity, mainCategory)
    # Some pages will not have these blocks
    if random.choice([True, False]):
        createSubHeaderFooter('subHeaderHeading', mainCity, mainCategory)
        createSubHeaderFooter('subHeaderParagraph', mainCity, mainCategory)

    createMainContainer('mainParagraph', mainCity, mainCategory)

    # Some pages will not have these blocks
    if random.choice([True, False]):
        createSubHeaderFooter('subFooterHeading', mainCity, mainCategory)
        createSubHeaderFooter('subFooterParagraph', mainCity, mainCategory)

    print('=> Content for ' + mainCity.upper() + ' ' + mainCategory.upper() + " created")


def filterTitle(mainCity, mainCategory):
    """
    This function returns correct title with name of city
    It combines mainCategory with titleDict keys and founds similar keys

    :param mainCity:
    :param mainCategory:
    :return: titleDict[mainCategory]
    """

    titleDict = {
        'design': '🥇 Web Design Agency in ' + mainCity + '. Web designers in ' + mainCity,
        'development': '🥇 Web Development Agency in ' + mainCity + '. Web developers in ' + mainCity,
        'magento': '🥇 Magento Web Development & eCommerce consulting agency in ' + mainCity,
        'shopify': '🥇 Shopify Development Agency in ' + mainCity + '. Web developers in ' + mainCity,
        'wordpress': '🥇 WordPress & WooCommerce Development Agency in ' + mainCity + '. Web developers in ' + mainCity
    }
    return titleDict[mainCategory]


def filterMeta(mainCity, mainCategory):
    """
    This function returns correct meta with name of city
    It combines mainCategory with metaDict keys and founds similar keys

    :param mainCity:
    :param mainCategory:
    :return: metaDict[mainCategory]
    """

    metaDict = {
        'design': 'Web design agency in ' + mainCity + ' ✅ with full-stack front-end back-end developers in ' + mainCity + '⚡',
        'development': 'Web development agency in ' + mainCity + ' ✅ with full-stack front-end back-end developers in ' + mainCity + '⚡',
        'magento': 'Magento agency in ' + mainCity + ' ✅ with certified developers and solution specialists ready to start today. ⚡We design, develop and support.',
        'shopify': 'Shopify agency in ' + mainCity + ' ✅ with full-stack front-end back-end developers in ' + mainCity + '.⚡',
        'wordpress': 'WordPress & WooCommerce agency in ' + mainCity + ' ✅ with full-stack front-end back-end developers in ' + mainCity + '. ⚡'
    }
    return metaDict[mainCategory]


def filterHeaderHeading(mainCity, mainCategory):
    """
    This function returns correct Header heading with name of city
    It combines mainCategory with headerHeadingDict keys and founds similar keys

    :param mainCity:
    :param mainCategory:
    :return: metaDict[mainCategory]
    """
    headerHeadingDict = {
        'design': 'Web design agency in <strong>' + mainCity + '</strong> with top-rated designers, developers, and marketing managers in <strong>' + mainCity + '</strong>',
        'development': 'Web development agency in <strong>' + mainCity + '</strong> with top-rated full-stack developers',
        'magento': 'Magento Agency in <strong>' + mainCity + '</strong>. Expert Magento Web Development in <strong>' + mainCity + '</strong>',
        'shopify': 'Shopify development agency in <strong>' + mainCity + '</strong> with top-rated full-stack developers',
        'wordpress': 'Wordpress &amp; WooCommerce development agency in <strong>' + mainCity + '</strong> with top-rated full-stack developers'
    }
    return headerHeadingDict[mainCategory]


def filterHeaderParagraph(mainCity, mainCategory):
    """
    This function returns correct Header paragraphs with name of city
    It combines mainCategory with headerParagraphDict keys and founds similar keys

    :param mainCity:
    :param mainCategory:
    :return: metaDict[mainCategory]
    """
    headerParagraphDict = {
        'design': 'We provide full-stack developers in ' + mainCity + '. Our agency support clients around Manchester and surrounding areas',
        'development': 'We provide full-stack development and support service in ' + mainCity + ' with a primary focus on month-by-month improvements to store resulting in better performance, rankings and revenue.',
        'magento': 'We provide full-stack Magento development and support service in ' + mainCity + ' with a primary focus on month-by-month improvements to store resulting in better performance, rankings and revenue.',
        'shopify': 'We provide full-stack Shopify development and support service ' + mainCity + ' with a primary focus on month-by-month improvements to store resulting in better performance, rankings and revenue.',
        'wordpress': 'We provide full-stack Wordpress & WooCommerce developers in ' + mainCity + '. Our WordPress & WooCommerce agency support clients around ' + mainCity + ' and surrounding areas'
    }
    return headerParagraphDict[mainCategory]


def createSubHeaderFooter(searchKey, mainCity, mainCategory):
    """
    This function create random 'underheader' and 'above-footer' content
    And then add it to resultDict
    """
    result = []
    # Add all founded values by searchKey to result[]
    for key, value in websitesDict.items():
        for subKey in websitesDict.get(key, {}):
            # Check if in websitesDict key with similar text exist
            if searchKey in str(subKey):
                cityToReplace = str(websitesDict[key]['City'])
                categoryToReplace = str(websitesDict[key]['Category'])
                contentToChange = str(websitesDict[key][subKey])

                # Additional keys
                for catKey, catValue in presets.categoryDict.items():
                    if str(presets.categoryDict[catKey]) in contentToChange:
                        contentToChange = contentToChange.replace(str(presets.categoryDict[catKey]), mainCategory)
                    else:
                        pass

                # Here we replace city and category in text
                contentToChange = contentToChange.replace(categoryToReplace, mainCategory.capitalize())
                contentToChange = contentToChange.replace(cityToReplace, mainCity)
                result.append(contentToChange)
            else:
                pass
    # Catch empty array
    try:
        result[0]
    except IndexError:
        pass
    else:
        # Get one random text from array without repeating
        resultContent = random.sample(result, 1)
        try:
            # We add searchKey to generate key 'mainSubHeader/Footer' with unique id
            resultDict['main' + searchKey] = resultContent[0]

        except IndexError:
            print('Index error when creating subHeader >' + mainCity.upper() + ' ' + mainCategory.upper())
            pass


def createMainContainer(searchKey, mainCity, mainCategory):
    """
    This function creates 'mainContainer' content
    1. We find all text what we needed
    2. Replace city and category
    3. Find correct heading to each category
    4. Add to resultDict

    :param searchKey:
    :param mainCity:
    :param mainCategory:
    """
    result = []
    i = 0
    mainContainerHeadings = {
        "design":
            {
                "Competition": 'Give Your Competition a <strong>Run for Its Money</strong>',
                "Design": 'The MageCloud Web <strong>Web Design Difference</strong>'
            },
        "development":
            {
                "Back-end": 'Back-end <strong>Development</strong>',
                "Front-end": 'Front-end <strong>Development</strong>',
                "Platform": "Platform <strong>Integrations</strong>",
                "Plugin": "Plugin <strong>Development</strong>"
            },
        "magento":
            {
                "Development": 'Your Magento Ecommerce <br> Development Partner in <strong>' + mainCity + '</strong>',
                "Design": 'Magento Ecommerce Design <br> in <strong>' + mainCity + '</strong>',
                "Marketing": 'Creative Marketing Strategy for <br> Magento in <strong>' + mainCity + '</strong>'
            },
        "shopify":
            {
                "Back-end": 'Back-end <strong>Development</strong>',
                "Front-end": 'Front-end <strong>Development</strong>',
                "Platform": "Platform <strong>Integrations</strong>",
                "Plugin": "Plugin <strong>Development</strong>"
            },
        "wordpress":
            {
                "Back-end": 'Back-end <strong>Development</strong>',
                "Front-end": 'Front-end <strong>Development</strong>',
                "Platform": "Platform <strong>Integrations</strong>",
                "Plugin": "Plugin <strong>Development</strong>"
            }
    }
    # Add all founded values by searchKey to result[]
    for key, value in websitesDict.items():
        # Here we need to get only mainParagraphs
        for subKey in websitesDict.get(key, {}):
            if searchKey in str(subKey):
                # Here we get current city name and category
                # Then check if this string is in paragraph
                # And replace to mainCity and mainCategory
                cityToReplace = str(websitesDict[key]['City'])
                categoryToReplace = str(websitesDict[key]['Category'])
                contentToChange = str(websitesDict[key][subKey])

                # Additional keys
                for catKey, catValue in presets.categoryDict.items():
                    if str(presets.categoryDict[catKey]) in contentToChange:
                        contentToChange = contentToChange.replace(str(presets.categoryDict[catKey]), mainCategory)
                    else:
                        pass
                # Here we edit all paragraph so we need to change city and category
                contentToChange = contentToChange.replace(cityToReplace, mainCity)
                contentToChange = contentToChange.replace(categoryToReplace, mainCategory.capitalize())

                result.append(contentToChange)
    # Catch empty a array
    try:
        result[0]
    except IndexError:
        pass
    else:
        # We need to check len of result list
        # And if len is more than 4 -> max number will be 4
        if len(result) > len(mainContainerHeadings[mainCategory]):
            resultParagraph = random.sample(result, 4)
        else:
            resultParagraph = random.sample(result, len(result))

        # Here we filter all data in loop
        # And add text to correct resultDict[key] + i (id for key)
        i = 0
        for key, value in mainContainerHeadings[mainCategory].items():
            # Here we check to which key we will add selected data
            # We are checking by searchKey
            try:
                resultDict['mainContainerHeading' + str(i)] = mainContainerHeadings[mainCategory][key]
                resultDict['mainContainerParagraph' + str(i)] = resultParagraph[i]
            # We need to check if size
            except IndexError:
                pass
                break
            else:
                i += 1


def saveFile(mainCity, mainCategory):
    """
    This function does last step of content creating

    1. In 'results' folder for every city we create folder
    2. Folder will have the same name with the city

    3. For every category file we add paragraphs and headings
    4. And then we save file with name "City-category"

    :param mainCity:
    :param mainCategory:
    :return:
    """

    # Create folder and file with specific category
    folder = 'results/' + mainCity
    file = mainCity + '-' + mainCategory + '.docx'
    try:
        os.mkdir(folder)
    except OSError:
        pass
    else:
        pass

    # Initialize document
    document = Document()

    # Add main headings, paragraphs, title and meta
    document.add_heading(mainCity + ' ' + mainCategory.capitalize(), 0)
    title = document.add_paragraph(resultDict['title'])
    meta = document.add_paragraph(resultDict['meta'])

    # From resultDict we get all content
    for key, value in resultDict.items():
        if 'Heading' in str(key):
            p = document.add_paragraph(resultDict[key])
        elif 'Paragraph' in str(key):
            p = document.add_paragraph(resultDict[key])

    document.add_page_break()
    # Save document in specific folder
    document.save(folder + '/' + file)
    print('==> ' + mainCity.upper() + ' pages have been created' + '\n')


if __name__ == '__main__':
    """
    Here we run two main functions
    contentGenerator() gets all data from links
    And creates websitesDict which contains structured info
    
    Then we run two loops (all cities, all categories)
    For every city we create 4 files (4 categories)
    """
    getFile('websites.txt', websitesList)
    getFile('cities.txt', citiesList)
    parseWebsites()
    for city in citiesList:
        for category in categoryList:
            contentCreator(city, category)
            saveFile(city, category)
